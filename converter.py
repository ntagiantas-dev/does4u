import streamlit as st
import docx
import PyPDF2
import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors

# ==========================================
# 🛠️ ΜΗΧΑΝΙΣΜΟΙ ΔΙΑΒΑΣΜΑΤΟΣ (ΑΠΟ ΤΟ ΑΡΧΕΙΟ ΣΟΥ)
# ==========================================

def read_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            # 8 κενά μέσα από το for
            content = page.extract_text()
            if content:
                # 12 κενά μέσα από το if
                text += content + "\n"
        # ΠΡΟΣΟΧΗ: Το return πρέπει να είναι στην ίδια ευθεία με το for!
        return text if text.strip() else "Δεν βρέθηκε αναγνώσιμο κείμενο."
    except Exception as e:
        # ΤΟ EXCEPT ΠΡΕΠΕΙ ΝΑ ΕΙΝΑΙ ΣΤΗΝ ΙΔΙΑ ΕΥΘΕΙΑ ΜΕ ΤΟ TRY
        return f"Σφάλμα ανάγνωσης PDF: {e}"

def read_docx(file):
    try:
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        return f"Σφάλμα ανάγνωσης Word: {e}"

def read_txt(file):
    try:
        return file.read().decode("utf-8")
    except:
        return file.read().decode("latin-1")

# ==========================================
# 📦 ΜΗΧΑΝΙΣΜΟΙ ΔΗΜΙΟΥΡΓΙΑΣ (ΓΙΑ ΤΟΝ ΠΕΛΑΤΗ)
# ==========================================

def create_pdf(text_content):
    """Μετατροπή σε PDF με αυτόματο κατέβασμα και υποστήριξη Ελληνικών"""
    import os
    import urllib.request
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    
    font_name = 'DejaVuSans'
    font_file = 'DejaVuSans.ttf'
    
    # Αν η γραμματοσειρά δεν υπάρχει τοπικά, τη κατεβάζουμε αυτόματα από το GitHub
    if not os.path.exists(font_file):
        url = "https://github.com/mcmaster-btech/mcmaster-btech.github.io/raw/master/fonts/DejaVuSans.ttf"
        try:
            urllib.request.urlretrieve(url, font_file)
        except Exception as e:
            pass

    # Προσπάθεια καταχώρησης της γραμματοσειράς
    try:
        pdfmetrics.registerFont(TTFont('DejaVuSans', font_file))
        font_name = 'DejaVuSans'
    except:
        font_name = 'Helvetica' # Fallback αν αποτύχει τελείως

    styles = getSampleStyleSheet()
    greek_style = ParagraphStyle('GStyle', parent=styles['Normal'], fontName=font_name, fontSize=11, leading=16)
    
    story = []
    lines = text_content.split('\n')
    for line in lines:
        if line.strip():
            story.append(Paragraph(line, greek_style))
            story.append(Spacer(1, 6))
            
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def create_docx(text_content):
    """Μετατροπή σε έγγραφο Word (.docx)"""
    doc = docx.Document()
    lines = text_content.split('\n')
    for line in lines:
        if line.strip():
            doc.add_paragraph(line)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ==========================================
# 🖥️ Η ΒΙΤΡΙΝΑ ΤΟΥ ΜΕΤΑΤΡΟΠΕΑ (UI)
# ==========================================
def show_converter_ui():
    st.write("---")
    st.markdown("### 🔄 Πολυμορφικός Μετατροπέας Εγγράφων")
    
    # Πιάνουμε το αρχείο από το 'conv_up' του app.py
    uploaded_file = st.session_state.get('conv_up', None)
    
    if uploaded_file is not None:
        file_name = uploaded_file.name
        file_type = file_name.split('.')[-1].lower()
        
        st.success(f"📥 Το αρχείο **{file_name}** φορτώθηκε επιτυχώς!")
        
        # 1. Πρώτα διαβάζουμε το περιεχόμενο του αρχείου που ανέβηκε
        raw_text = ""
        if file_type == "pdf":
            raw_text = read_pdf(uploaded_file)
        elif file_type == "docx":
            raw_text = read_docx(uploaded_file)
        elif file_type == "txt":
            raw_text = read_txt(uploaded_file)
        else:
            st.warning("⚠️ Για φωτογραφίες (.jpg, .png) απαιτείται σύνδεση με OCR.")
            return

        # 2. ΡΩΤΑΜΕ ΤΟΝ ΧΡΗΣΤΗ: Τι μετατροπή θέλει;
        st.markdown("#### 🎯 Επιλέξτε τη μορφή του τελικού εγγράφου:")
        
        # Η λίστα με τις διαθέσιμες μετατροπές
        conversion_target = st.selectbox(
            "Διαθέσιμες Μετατροπές:",
            [
                "--- Παρακαλώ επιλέξτε ---",
                "Μετατροπή σε Αρχείο Word (.docx)",
                "Μετατροπή σε Έγγραφο PDF (.pdf)",
                "Μετατροπή σε Απλό Κείμενο (.txt)"
            ],
            key=f"selectbox_converter_unique_fixed_id_{key_suffix}"
        )
        
        # 3. Ετοιμάζουμε το αρχείο ανάλογα με την επιλογή του
        if conversion_target == "Μετατροπή σε Αρχείο Word (.docx)":
            with st.spinner("Ετοιμάζω το Word..."):
                docx_bytes = create_docx(raw_text)
                st.download_button(
                    label="📥 Λήψη έτοιμου εγγράφου Word (.docx)",
                    data=docx_bytes,
                    file_name=f"does4u_converted.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
        elif conversion_target == "Μετατροπή σε Έγγραφο PDF (.pdf)":
            with st.spinner("Ετοιμάζω το PDF..."):
                pdf_bytes = create_pdf(raw_text)
                st.download_button(
                    label="📥 Λήψη έτοιμου εγγράφου PDF (.pdf)",
                    data=pdf_bytes,
                    file_name=f"does4u_converted.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
                
        elif conversion_target == "Μετατροπή σε Απλό Κείμενο (.txt)":
            st.download_button(
                label="📥 Λήψη αρχείου Κειμένου (.txt)",
                data=raw_text,
                file_name=f"does4u_converted.txt",
                mime="text/plain",
                use_container_width=True
            )
            
    else:
        st.warning("⚠️ Παρακαλώ ανεβάστε πρώτα ένα έγγραφο στο δεξί κουτί (Convert Box).")